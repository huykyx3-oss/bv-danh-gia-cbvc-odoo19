import io
import json
from odoo import http
from odoo.http import request, content_disposition


class ExportDocxController(http.Controller):

    @http.route('/bv_danh_gia/export_mau01/<int:eval_id>', type='http', auth='user')
    def export_mau01_docx(self, eval_id, **kwargs):
        """Export Mẫu 01 - Phiếu theo dõi, đánh giá CBVC as DOCX."""
        evaluation = request.env['bv.monthly.evaluation'].browse(eval_id)
        if not evaluation.exists():
            return request.not_found()

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            return request.make_response(
                json.dumps({'error': 'Thư viện python-docx chưa được cài đặt. '
                            'Chạy: pip install python-docx'}),
                headers=[('Content-Type', 'application/json')])

        def _fmt_score(value):
            """Format score with Vietnamese decimal comma."""
            try:
                num = float(value or 0.0)
            except Exception:
                return ''
            if abs(num - round(num)) < 1e-9:
                return str(int(round(num)))
            # keep up to 2 decimals, trim trailing zero, then convert dot->comma
            txt = f'{num:.2f}'.rstrip('0').rstrip('.')
            return txt.replace('.', ',')

        # ------------------------------------------------------------------ #
        # Document base style
        # ------------------------------------------------------------------ #
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)

        # ------------------------------------------------------------------ #
        # Nhãn "Mẫu số 01" — canh phải, trên cùng
        # ------------------------------------------------------------------ #
        p_mau = doc.add_paragraph()
        p_mau.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_mau = p_mau.add_run('Mẫu số 01')
        run_mau.italic = True
        run_mau.font.size = Pt(12)

        # ------------------------------------------------------------------ #
        # Header — bảng 2 cột, không border
        # ------------------------------------------------------------------ #
        table_header = doc.add_table(rows=1, cols=2)
        table_header.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell_left = table_header.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_left.add_run('SỞ Y TẾ TỈNH QUẢNG NINH')
        run.bold = True
        run.font.size = Pt(12)
        run2 = p_left.add_run('\nBỆNH VIỆN ĐA KHOA TỈNH')
        run2.bold = True
        run2.font.size = Pt(12)

        cell_right = table_header.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_right.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
        run.bold = True
        run.font.size = Pt(12)
        run2 = p_right.add_run('\nĐộc lập – Tự do – Hạnh phúc')
        run2.bold = True
        run2.font.size = Pt(12)
        p_right.add_run('\n──────────────').font.size = Pt(12)

        doc.add_paragraph()

        # ------------------------------------------------------------------ #
        # Tiêu đề
        # ------------------------------------------------------------------ #
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('PHIẾU THEO DÕI, ĐÁNH GIÁ CÁN BỘ, CÔNG CHỨC, VIÊN CHỨC')
        run.bold = True
        run.font.size = Pt(14)

        month_labels = {
            '1': 'Tháng 1', '2': 'Tháng 2', '3': 'Tháng 3',
            '4': 'Tháng 4', '5': 'Tháng 5', '6': 'Tháng 6',
            '7': 'Tháng 7', '8': 'Tháng 8', '9': 'Tháng 9',
            '10': 'Tháng 10', '11': 'Tháng 11', '12': 'Tháng 12',
        }
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(
            f'(Kỳ theo dõi, đánh giá: '
            f'{month_labels.get(str(evaluation.month), str(evaluation.month))} '
            f'năm {evaluation.year})'
        ).font.size = Pt(12)

        # ------------------------------------------------------------------ #
        # Thông tin cá nhân
        # ------------------------------------------------------------------ #
        doc.add_paragraph(f'Họ và tên: {evaluation.employee_id.name or ""}')
        doc.add_paragraph(f'Chức vụ, chức danh: {evaluation.job_id.name or ""}')
        doc.add_paragraph(f'Đơn vị công tác: {evaluation.department_id.name or ""}')

        # ------------------------------------------------------------------ #
        # I. Bảng tiêu chí chung — 5 cột
        # ------------------------------------------------------------------ #
        h1 = doc.add_paragraph()
        run = h1.add_run('I. KẾT QUẢ THEO DÕI, ĐÁNH GIÁ THEO TIÊU CHÍ CHUNG')
        run.bold = True

        # Số La Mã cho nhóm cha
        ROMAN = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        col_headers = [
            'TT', 'Tiêu chí chấm điểm', 'Điểm tối đa',
            'Điểm tự chấm', 'Trưởng khoa/phòng chấm',
        ]
        for i, h in enumerate(col_headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

        # Dòng chỉ số cột theo mẫu: (1) (2) (3)
        index_row = table.add_row()
        index_row.cells[0].text = '(1)'
        index_row.cells[1].text = '(2)'
        index_row.cells[2].text = '(3)'
        index_row.cells[3].text = ''
        index_row.cells[4].text = ''
        for i in [0, 1, 2, 3, 4]:
            for p in index_row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

        # Nhóm các line theo parent_criteria_id (giữ thứ tự xuất hiện)
        # parent_criteria_id là related field: criteria_id.parent_id
        seen_parents = []
        groups = {}  # parent_id (hoặc False) -> [lines]
        for line in evaluation.criteria_line_ids:
            parent = line.parent_criteria_id  # Many2one record hoặc False
            parent_key = parent.id if parent else False
            if parent_key not in groups:
                groups[parent_key] = {'parent': parent, 'lines': []}
                seen_parents.append(parent_key)
            groups[parent_key]['lines'].append(line)

        for group_idx, parent_key in enumerate(seen_parents):
            group = groups[parent_key]
            parent = group['parent']
            group_lines = group['lines']

            # Dòng nhóm (I / II / III ...) — nếu có parent
            if parent:
                group_row = table.add_row()
                roman = ROMAN[group_idx] if group_idx < len(ROMAN) else str(group_idx + 1)
                group_row.cells[0].text = roman
                group_row.cells[1].text = parent.name or ''
                # Lấy điểm nhóm theo tổng các dòng con để khớp cấu trúc mẫu
                group_row.cells[2].text = _fmt_score(sum(l.max_score for l in group_lines))
                group_row.cells[3].text = ''
                group_row.cells[4].text = ''
                for i in [0, 2, 3, 4]:
                    for p in group_row.cells[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for c in range(5):
                    for p in group_row.cells[c].paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(11)

            # Dòng con
            sub_counter = 0
            for line in group_lines:
                sub_counter += 1
                row = table.add_row()
                row.cells[0].text = str(sub_counter)
                row.cells[1].text = line.criteria_id.name or ''
                row.cells[2].text = _fmt_score(line.max_score)
                row.cells[3].text = _fmt_score(line.self_score)
                row.cells[4].text = _fmt_score(line.dept_score)
                for i in [0, 2, 3, 4]:
                    for p in row.cells[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for p in row.cells[1].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(11)

        # Dòng tổng cộng — merge 2 ô đầu
        total_row = table.add_row()
        merged_total = total_row.cells[0].merge(total_row.cells[1])
        p_tc = merged_total.paragraphs[0]
        p_tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tc = p_tc.add_run('Tổng cộng')
        run_tc.bold = True
        run_tc.font.size = Pt(11)

        general_max = evaluation.general_score_max or 30.0
        total_row.cells[2].text = _fmt_score(general_max)
        total_row.cells[3].text = _fmt_score(sum(evaluation.criteria_line_ids.mapped('self_score')))
        total_row.cells[4].text = _fmt_score(sum(evaluation.criteria_line_ids.mapped('dept_score')))
        for i in [2, 3, 4]:
            for p in total_row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

        # ------------------------------------------------------------------ #
        # II. Tổng hợp kết quả
        # ------------------------------------------------------------------ #
        doc.add_paragraph()
        h2 = doc.add_paragraph()
        run = h2.add_run(
            'II. TỔNG HỢP KẾT QUẢ THEO DÕI, ĐÁNH GIÁ CÁN BỘ, CÔNG CHỨC, VIÊN CHỨC'
        )
        run.bold = True

        doc.add_paragraph(f'1. Điểm tiêu chí chung: {_fmt_score(evaluation.general_score)}')
        doc.add_paragraph('2. Điểm tiêu chí kết quả thực hiện nhiệm vụ:')

        def _fmt_evidence_name(field_prefix):
            return getattr(evaluation, field_prefix + '_name', '') or ''

        basic_criteria = [
            ('a', 'số lượng kết quả thực hiện nhiệm vụ',
             'pct_quantity', 'dept_pct_quantity', 'evidence_quantity'),
            ('b', 'chất lượng kết quả thực hiện nhiệm vụ',
             'pct_quality', 'dept_pct_quality', 'evidence_quality'),
            ('c', 'tiến độ kết quả thực hiện nhiệm vụ',
             'pct_progress', 'dept_pct_progress', 'evidence_progress'),
        ]
        manager_criteria = [
            ('d', 'kết quả hoạt động của lĩnh vực được giao lãnh đạo, quản lý, phụ trách',
             'pct_field_result', 'dept_pct_field_result', 'evidence_field_result'),
            ('đ', 'khả năng tổ chức triển khai thực hiện nhiệm vụ',
             'pct_organization', 'dept_pct_organization', 'evidence_organization'),
            ('e', 'năng lực tập hợp, đoàn kết công chức thuộc phạm vi quản lý',
             'pct_team_cohesion', 'dept_pct_team_cohesion', 'evidence_team_cohesion'),
        ]
        all_criteria = basic_criteria + (manager_criteria if evaluation.is_manager else [])
        for lbl, desc, nv_f, tp_f, ev_f in all_criteria:
            nv_val = _fmt_score(getattr(evaluation, nv_f, 0) or 0)
            tp_val = getattr(evaluation, tp_f, 0) or 0
            ev_name = _fmt_evidence_name(ev_f)
            tp_part = f'; TP chấm: {_fmt_score(tp_val)}%' if tp_val else ''
            ev_part = f'; Minh chứng: {ev_name}' if ev_name else ''
            doc.add_paragraph(
                f'   - {lbl} là điểm tỷ lệ phần trăm (%) về {desc}: '
                f'NV tự chấm {nv_val}%{tp_part}{ev_part}'
            )

        role_label = 'giữ chức vụ' if evaluation.is_manager else 'không giữ chức vụ'
        doc.add_paragraph(
            f'   Điểm tiêu chí kết quả thực hiện nhiệm vụ '
            f'(đối với CBCCVC {role_label}): '
            f'{_fmt_score(evaluation.task_score)} điểm'
        )

        p_total = doc.add_paragraph()
        run = p_total.add_run(
            f'3. Tổng điểm theo dõi, đánh giá cán bộ, công chức, viên chức: '
            f'{_fmt_score(evaluation.total_score)} điểm'
        )
        run.bold = True

        doc.add_paragraph(f'4. Ưu điểm: {evaluation.strengths or ""}')
        doc.add_paragraph(f'5. Hạn chế, khuyết điểm: {evaluation.weaknesses or ""}')
        doc.add_paragraph(
            f'6. Ý kiến nhận xét của cấp có thẩm quyền theo dõi, đánh giá: '
            f'{evaluation.authority_comment or ""}'
        )

        # ------------------------------------------------------------------ #
        # Chữ ký — bảng 1 hàng, 2 cột (gộp ngày + chức danh + họ tên vào
        # cùng một ô, đúng theo mẫu docx)
        # ------------------------------------------------------------------ #
        doc.add_paragraph()
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        sig_data = [
            ('CBCCVC TỰ ĐÁNH GIÁ', evaluation.employee_id.name or ''),
            ('CẤP CÓ THẨM QUYỀN\nTHEO DÕI, ĐÁNH GIÁ', ''),
        ]
        for col_idx, (role_title, signer_name) in enumerate(sig_data):
            cell = sig_table.cell(0, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_date = p.add_run('....., ngày.... tháng..... năm.....')
            run_date.italic = True
            run_date.font.size = Pt(12)

            p.add_run('\n')
            run_title = p.add_run(role_title)
            run_title.bold = True
            run_title.font.size = Pt(12)

            p.add_run('\n')
            run_sub = p.add_run('(Ký tên, ghi rõ họ tên)')
            run_sub.italic = True
            run_sub.font.size = Pt(12)

            if signer_name:
                p.add_run(f'\n\n\n\n{signer_name}').font.size = Pt(12)

        # ------------------------------------------------------------------ #
        # Xuất ra buffer
        # ------------------------------------------------------------------ #
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = (
            f'Mau_01_{evaluation.employee_id.name}_'
            f'{evaluation.month}_{evaluation.year}.docx'
        )
        return request.make_response(
            buffer.getvalue(),
            headers=[
                ('Content-Type',
                 'application/vnd.openxmlformats-officedocument'
                 '.wordprocessingml.document'),
                ('Content-Disposition', content_disposition(filename)),
            ]
        )

    @http.route('/bv_danh_gia/export_mau02/<int:record_id>', type='http', auth='user')
    def export_mau02_docx(self, record_id, **kwargs):
        """Export Mẫu 02 - Phiếu xếp loại chất lượng CBVC as DOCX."""
        record = request.env['bv.yearly.classification'].browse(record_id)
        if not record.exists():
            return request.not_found()

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            return request.make_response(
                json.dumps({'error': 'Thư viện python-docx chưa được cài đặt.'}),
                headers=[('Content-Type', 'application/json')])

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)

        # Header
        table_header = doc.add_table(rows=1, cols=2)
        cell_left = table_header.cell(0, 0)
        cell_right = table_header.cell(0, 1)

        p = cell_left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('TÊN CƠ QUAN,\nTỔ CHỨC, ĐƠN VỊ\n───────')
        r.bold = True
        r.font.size = Pt(12)

        p = cell_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n───────────────')
        r.bold = True
        r.font.size = Pt(12)

        doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run('PHIẾU XẾP LOẠI CHẤT LƯỢNG CÁN BỘ, CÔNG CHỨC, VIÊN CHỨC')
        r.bold = True
        r.font.size = Pt(14)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f'Năm {record.year}')

        doc.add_paragraph(f'Họ và tên: {record.employee_id.name or ""}')
        doc.add_paragraph(f'Chức vụ, chức danh: {record.job_id.name or ""}')
        doc.add_paragraph(f'Đơn vị công tác: {record.department_id.name or ""}')

        # I. Tổng hợp
        h1 = doc.add_paragraph()
        h1.add_run('I. TỔNG HỢP KẾT QUẢ THEO DÕI, ĐÁNH GIÁ').bold = True

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['Tháng', 'Tổng điểm', 'Ghi chú']):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True

        month_data = [
            (1, record.month_1), (2, record.month_2), (3, record.month_3),
            ('Quý I', record.quarter_1_avg, 'TB cộng 3 tháng'),
            (4, record.month_4), (5, record.month_5), (6, record.month_6),
            ('Quý II', record.quarter_2_avg, 'TB cộng 3 tháng'),
            (7, record.month_7), (8, record.month_8), (9, record.month_9),
            ('Quý III', record.quarter_3_avg, 'TB cộng 3 tháng'),
            (10, record.month_10), (11, record.month_11), (12, record.month_12),
            ('Quý IV', record.quarter_4_avg, 'TB cộng 3 tháng'),
            ('TB cả năm', record.yearly_average, 'TB cộng 12 tháng'),
        ]

        for item in month_data:
            row = table.add_row()
            if len(item) == 3:
                row.cells[0].text = str(item[0])
                row.cells[1].text = f'{item[1]:.2f}'
                row.cells[2].text = item[2]
                for p in row.cells[0].paragraphs:
                    for r in p.runs:
                        r.bold = True
            else:
                row.cells[0].text = str(item[0])
                row.cells[1].text = str(item[1]) if item[1] > 0 else ''
                row.cells[2].text = ''
            for i in [0, 1]:
                for p in row.cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        clf_labels = {
            'excellent': 'Hoàn thành xuất sắc nhiệm vụ',
            'good': 'Hoàn thành tốt nhiệm vụ',
            'fair': 'Hoàn thành nhiệm vụ',
            'poor': 'Không hoàn thành nhiệm vụ',
        }

        # II. Tự xếp loại
        doc.add_paragraph()
        h2 = doc.add_paragraph()
        h2.add_run('II. CBCCVC TỰ ĐÁNH GIÁ, XẾP LOẠI CHẤT LƯỢNG').bold = True
        for clf_key, clf_label in clf_labels.items():
            marker = '■' if record.self_classification == clf_key else '□'
            doc.add_paragraph(f'   {marker} {clf_label}')

        # III. Kết quả
        h3 = doc.add_paragraph()
        h3.add_run('III. KẾT QUẢ XẾP LOẠI CỦA CẤP CÓ THẨM QUYỀN').bold = True
        for clf_key, clf_label in clf_labels.items():
            marker = '■' if record.final_classification == clf_key else '□'
            doc.add_paragraph(f'   {marker} {clf_label}')

        # IV. Đề xuất
        h4 = doc.add_paragraph()
        h4.add_run('IV. ĐỀ XUẤT PHƯƠNG ÁN XỬ LÝ').bold = True
        doc.add_paragraph(record.proposed_action or '')

        # Signatures
        doc.add_paragraph()
        sig_table = doc.add_table(rows=4, cols=2)
        for col_idx, (t, s) in enumerate([
            ('CBCCVC TỰ XẾP LOẠI', '(Ký tên, ghi rõ họ tên)'),
            ('CẤP CÓ THẨM QUYỀN XẾP LOẠI', '(Ký tên, ghi rõ họ tên)'),
        ]):
            sig_table.cell(0, col_idx).text = '......, ngày .... tháng .... năm ....'
            for p in sig_table.cell(0, col_idx).paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = sig_table.cell(1, col_idx).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(t).bold = True
            p = sig_table.cell(2, col_idx).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(s).italic = True

        p = sig_table.cell(3, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'\n\n\n{record.employee_id.name or ""}')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'Mau_02_{record.employee_id.name}_{record.year}.docx'
        return request.make_response(
            buffer.getvalue(),
            headers=[
                ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                ('Content-Disposition', content_disposition(filename)),
            ])